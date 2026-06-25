"""
Document processor for HarveySpecter.

Extracts text with page/line-level metadata from PDF, DOCX, plain-text,
and image files.  Every extraction path returns a unified ``DocumentResult``
so downstream services never need to care about the original file format.
"""

from __future__ import annotations

import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.models.schemas import DocumentResult, PageContent, PageLine

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════
#  Public API
# ═══════════════════════════════════════════════════════════════════════


async def process_document(file_path: str) -> DocumentResult:
    """Route to the correct extractor based on file extension.

    Args:
        file_path: Absolute or relative path to the uploaded file.

    Returns:
        A ``DocumentResult`` with full text, per-page line data, and metadata.

    Raises:
        ValueError: If the file type is not supported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return await extract_pdf(file_path)
    elif suffix == ".docx":
        return await extract_docx(file_path)
    elif suffix in (".txt", ".text", ".md"):
        text = path.read_text(encoding="utf-8")
        return await extract_text(text)
    elif suffix in (".png", ".jpg", ".jpeg", ".tiff", ".tif"):
        return await extract_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# ═══════════════════════════════════════════════════════════════════════
#  PDF — PyMuPDF (fitz)
# ═══════════════════════════════════════════════════════════════════════


async def extract_pdf(file_path: str) -> DocumentResult:
    """Extract text from a PDF, preserving page and line numbers.

    Uses ``page.get_text("words")`` to obtain word-level bounding boxes,
    then groups words by their *block_no* and *line_no* fields to
    reconstruct ordered lines with positional metadata.

    Args:
        file_path: Path to a PDF file.

    Returns:
        ``DocumentResult`` populated with page/line data.
    """
    doc = fitz.open(file_path)
    pages: list[PageContent] = []
    full_text_parts: list[str] = []
    global_line_counter = 0

    try:
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_number = page_idx + 1

            # get_text("words") → list of (x0, y0, x1, y1, word, block_no, line_no, word_no)
            words = page.get_text("words")

            # Group words by (block_no, line_no) to reconstruct lines
            line_groups: dict[tuple[int, int], list[str]] = {}
            for word_info in words:
                word_text = word_info[4]
                block_no = word_info[5]
                line_no = word_info[6]
                key = (block_no, line_no)
                if key not in line_groups:
                    line_groups[key] = []
                line_groups[key].append(word_text)

            # Sort by block_no then line_no to maintain reading order
            sorted_keys = sorted(line_groups.keys())

            page_lines: list[PageLine] = []
            page_text_parts: list[str] = []

            for key in sorted_keys:
                global_line_counter += 1
                line_text = " ".join(line_groups[key])
                page_lines.append(
                    PageLine(line_number=global_line_counter, text=line_text)
                )
                page_text_parts.append(line_text)

            page_raw = "\n".join(page_text_parts)
            pages.append(
                PageContent(
                    page_number=page_number,
                    lines=page_lines,
                    raw_text=page_raw,
                )
            )
            full_text_parts.append(page_raw)

        full_text = "\n\n".join(full_text_parts)

        return DocumentResult(
            text=full_text,
            pages=pages,
            metadata={
                "file_type": "pdf",
                "page_count": len(doc),
                "file_path": file_path,
                "total_lines": global_line_counter,
            },
        )
    finally:
        doc.close()


# ═══════════════════════════════════════════════════════════════════════
#  DOCX — python-docx
# ═══════════════════════════════════════════════════════════════════════


async def extract_docx(file_path: str) -> DocumentResult:
    """Extract text from a DOCX file with paragraph-level tracking.

    Heading styles (``Heading 1``, ``Heading 2``, …) are preserved in the
    metadata so the knowledge-base ingestion pipeline can use them for
    structure-aware chunking.

    Args:
        file_path: Path to a ``.docx`` file.

    Returns:
        ``DocumentResult`` with paragraph ≈ line mapping.
    """
    doc = DocxDocument(file_path)

    pages: list[PageContent] = []
    full_text_parts: list[str] = []
    current_page_lines: list[PageLine] = []
    current_page_texts: list[str] = []
    line_counter = 0
    page_number = 1

    # DOCX doesn't have native "pages", but we create virtual pages
    # every ~50 paragraphs or on Heading 1 boundaries.
    PARAGRAPHS_PER_PAGE = 50

    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        line_counter += 1
        style_name = paragraph.style.name if paragraph.style else ""

        # Prefix headings for visibility
        if style_name.startswith("Heading"):
            heading_prefix = f"[{style_name}] "
            text = heading_prefix + text

        current_page_lines.append(PageLine(line_number=line_counter, text=text))
        current_page_texts.append(text)

        # Virtual page break on heading-1 or every N paragraphs
        is_page_break = (
            style_name == "Heading 1"
            or len(current_page_lines) >= PARAGRAPHS_PER_PAGE
        )

        if is_page_break and len(current_page_lines) > 1:
            raw = "\n".join(current_page_texts)
            pages.append(
                PageContent(
                    page_number=page_number,
                    lines=current_page_lines,
                    raw_text=raw,
                )
            )
            full_text_parts.append(raw)
            page_number += 1
            current_page_lines = []
            current_page_texts = []

    # Flush remaining lines
    if current_page_lines:
        raw = "\n".join(current_page_texts)
        pages.append(
            PageContent(
                page_number=page_number,
                lines=current_page_lines,
                raw_text=raw,
            )
        )
        full_text_parts.append(raw)

    full_text = "\n\n".join(full_text_parts)

    return DocumentResult(
        text=full_text,
        pages=pages,
        metadata={
            "file_type": "docx",
            "page_count": len(pages),
            "file_path": file_path,
            "total_lines": line_counter,
            "total_paragraphs": len(doc.paragraphs),
        },
    )


# ═══════════════════════════════════════════════════════════════════════
#  Plain text
# ═══════════════════════════════════════════════════════════════════════


def extract_text_sync(content: str) -> DocumentResult:
    """Wrap raw text in a ``DocumentResult`` with line numbering.

    Args:
        content: Raw text content (already read from file or pasted by user).

    Returns:
        ``DocumentResult`` with a single virtual page.
    """
    raw_lines = content.splitlines()
    page_lines: list[PageLine] = []

    for idx, line in enumerate(raw_lines, start=1):
        page_lines.append(PageLine(line_number=idx, text=line))

    page = PageContent(
        page_number=1,
        lines=page_lines,
        raw_text=content,
    )

    return DocumentResult(
        text=content,
        pages=[page],
        metadata={
            "file_type": "text",
            "page_count": 1,
            "total_lines": len(raw_lines),
        },
    )


async def extract_text(content: str) -> DocumentResult:
    """Wrap raw text in a ``DocumentResult`` with line numbering.

    Args:
        content: Raw text content (already read from file or pasted by user).

    Returns:
        ``DocumentResult`` with a single virtual page.
    """
    return extract_text_sync(content)


# ═══════════════════════════════════════════════════════════════════════
#  Image / OCR (placeholder)
# ═══════════════════════════════════════════════════════════════════════


async def extract_image(file_path: str) -> DocumentResult:
    """Placeholder for image-based document extraction via OCR.

    A production implementation would use Google Document AI or
    Tesseract to perform OCR.  For the MVP, this returns a helpful
    message indicating that OCR is not yet configured.

    Args:
        file_path: Path to an image file (PNG, JPEG, TIFF).

    Returns:
        ``DocumentResult`` with a placeholder message.
    """
    logger.warning("OCR extraction requested but not yet implemented: %s", file_path)

    placeholder_text = (
        "[OCR extraction is not yet configured. "
        "To enable image-based document processing, integrate Google Document AI "
        "or Tesseract OCR and implement the extract_image() function in "
        "app/services/document_processor.py]"
    )

    page = PageContent(
        page_number=1,
        lines=[PageLine(line_number=1, text=placeholder_text)],
        raw_text=placeholder_text,
    )

    return DocumentResult(
        text=placeholder_text,
        pages=[page],
        metadata={
            "file_type": "image",
            "page_count": 1,
            "total_lines": 1,
            "file_path": file_path,
            "ocr_enabled": False,
        },
    )


class DocumentProcessor:
    """Service class wrapper for document text extraction."""

    def process_document(self, file_path: str) -> DocumentResult:
        import asyncio
        return asyncio.run(process_document(file_path))

    def extract_pdf(self, file_path: str) -> DocumentResult:
        import asyncio
        return asyncio.run(extract_pdf(file_path))

    def extract_docx(self, file_path: str) -> DocumentResult:
        import asyncio
        return asyncio.run(extract_docx(file_path))

    def extract_text(self, text: str) -> DocumentResult:
        return extract_text_sync(text)

    def extract_image(self, file_path: str) -> DocumentResult:
        import asyncio
        return asyncio.run(extract_image(file_path))

